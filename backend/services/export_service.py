import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

class ExportService:
    def __init__(self):
        self.export_dir = "exports"
        os.makedirs(self.export_dir, exist_ok=True)

    async def export_meeting(self, meeting, fmt):
        if fmt == "txt":
            return await self._export_txt(meeting)
        elif fmt == "md":
            return await self._export_md(meeting)
        elif fmt == "docx":
            return await self._export_docx(meeting)
        elif fmt == "pdf":
            return await self._export_pdf(meeting)
        else:
            raise ValueError(f"Unsupported export format: {fmt}")

    async def _export_txt(self, meeting):
        path = os.path.join(self.export_dir, f"{meeting.id}.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write(meeting.summary or "")
        return path

    async def _export_md(self, meeting):
        path = os.path.join(self.export_dir, f"{meeting.id}.md")
        with open(path, "w", encoding="utf-8") as f:
            f.write("# Meeting Summary\n")
            f.write(meeting.summary or "")
        return path

    async def _export_docx(self, meeting):
        path = os.path.join(self.export_dir, f"{meeting.id}.docx")
        doc = Document()
        doc.add_heading("Meeting Summary", 0)
        doc.add_paragraph(meeting.summary or "")
        if meeting.action_items:
            doc.add_heading("Action Items", level=1)
            for item in meeting.action_items:
                doc.add_paragraph(item, style="List Bullet")
        doc.save(path)
        return path

    async def _export_pdf(self, meeting):
        path = os.path.join(self.export_dir, f"{meeting.id}.pdf")
        c = canvas.Canvas(path, pagesize=letter)
        text_obj = c.beginText(50, 750)
        text_obj.setFont("Helvetica", 12)

        text_obj.textLine("Meeting Summary")
        text_obj.textLine("")
        text_obj.textLines(meeting.summary or "")

        if meeting.action_items:
            text_obj.textLine("")
            text_obj.textLine("Action Items:")
            for item in meeting.action_items:
                text_obj.textLine(f"- {item}")

        if meeting.decisions:
            text_obj.textLine("")
            text_obj.textLine("Decisions:")
            for d in meeting.decisions:
                text_obj.textLine(f"- {d}")

        if meeting.participants:
            text_obj.textLine("")
            text_obj.textLine("Participants: " + ", ".join(meeting.participants))

        if meeting.duration:
            text_obj.textLine("")
            text_obj.textLine(f"Duration: {meeting.duration}")

        c.drawText(text_obj)
        c.save()
        return path
